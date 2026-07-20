"""
Resume upload parsing (Phase 3).

Extracts structured resume data from an uploaded PDF or DOCX so the
student lands on a pre-filled, editable form instead of a blank one.

Pipeline:
    1. Text extraction — PyMuPDF for PDFs, python-docx for DOCX.
       OCR fallback (Tesseract) when a PDF page has no text layer
       (scanned/image resumes).
    2. Section detection — split lines into sections by matching headers
       against a whitelist with fuzzy matching (rapidfuzz).
    3. Field extraction — contact info via regex, education / experience /
       projects via date-range + spaCy NER heuristics, skills via
       delimiter splitting.
    4. Confidence — every extracted field gets "high" or "low" so the UI
       can highlight what needs verification (green vs yellow).

Rule-based only — no LLM involved.
"""

from __future__ import annotations

import io
import logging
import re
from typing import Any

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Optional dependencies (graceful degradation, same pattern as bullet_analyzer)
# ---------------------------------------------------------------------------

_SPACY_NLP = None
try:
    import spacy as _spacy_mod

    try:
        _SPACY_NLP = _spacy_mod.load("en_core_web_sm")
    except Exception:
        logger.debug("spaCy model not installed; NER extraction disabled.")
except ImportError:
    pass

_OCR_AVAILABLE = False
try:
    import pytesseract  # noqa: F401
    from PIL import Image  # noqa: F401

    try:
        pytesseract.get_tesseract_version()
        _OCR_AVAILABLE = True
    except Exception:
        logger.debug("Tesseract binary not found; OCR fallback disabled.")
except ImportError:
    pass

try:
    from rapidfuzz import fuzz as _fuzz
except ImportError:
    _fuzz = None

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

MAX_FILE_SIZE = 5 * 1024 * 1024  # 5 MB
ALLOWED_EXTENSIONS = {".pdf", ".docx"}

# Canonical section -> header aliases (lowercase). Matched fuzzily so
# "PROFESSIONAL EXPERIENCE" or "Educational Background" still land right.
_SECTION_ALIASES: dict[str, list[str]] = {
    "summary": [
        "summary", "professional summary", "objective", "career objective",
        "profile", "about me", "about",
    ],
    "experience": [
        "experience", "work experience", "professional experience",
        "employment", "employment history", "work history", "internships",
        "internship experience",
    ],
    "education": [
        "education", "educational background", "academic background",
        "academics", "qualifications", "academic qualifications",
    ],
    "projects": [
        "projects", "personal projects", "academic projects",
        "key projects", "project work",
    ],
    "skills": [
        "skills", "technical skills", "skills & tools", "core competencies",
        "technologies", "tech stack", "skills and abilities",
    ],
    "certifications": [
        "certifications", "certificates", "certifications & courses",
        "courses", "licenses", "achievements & certifications",
    ],
}

_EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
_PHONE_RE = re.compile(r"(?:\+91[\s-]?)?[6-9]\d{4}[\s-]?\d{5}\b")
_LINKEDIN_RE = re.compile(r"(?:https?://)?(?:www\.)?linkedin\.com/in/[\w%-]+/?", re.I)
_URL_RE = re.compile(r"https?://\S+|(?:www\.)\S+\.\S+", re.I)

# "Jan 2022 - Mar 2023", "2021-2023", "June 2022 – Present" ...
_MONTHS = (
    "jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|"
    "jul(?:y)?|aug(?:ust)?|sep(?:t(?:ember)?)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?"
)
_DATE_TOKEN = rf"(?:(?:{_MONTHS})[a-z]*\.?\s*)?(?:19|20)\d{{2}}"
_DATE_RANGE_RE = re.compile(
    rf"({_DATE_TOKEN})\s*(?:-|–|—|to)\s*({_DATE_TOKEN}|present|current|now|ongoing)",
    re.I,
)

_DEGREE_RE = re.compile(
    r"\b(b\.?\s?tech|m\.?\s?tech|b\.?\s?e\.?|m\.?\s?e\.?|b\.?\s?sc|m\.?\s?sc|"
    r"bca|mca|bba|mba|b\.?\s?com|m\.?\s?com|ph\.?d|bachelor(?:'?s)?|master(?:'?s)?|diploma)\b",
    re.I,
)

_GPA_RE = re.compile(
    r"\b(?:cgpa|gpa|percentage|score)\s*[:\-]?\s*(\d{1,2}(?:\.\d{1,2})?\s*/?\s*(?:10|100|4)?|\d{1,3}\s*%)",
    re.I,
)

_MONTH_NUM = {
    "jan": "01", "feb": "02", "mar": "03", "apr": "04", "may": "05",
    "jun": "06", "jul": "07", "aug": "08", "sep": "09", "oct": "10",
    "nov": "11", "dec": "12",
}

_SKILL_SPLIT_RE = re.compile(r"[,|•▪‣·;/]+")

# Fuzzy-match threshold for section headers
_HEADER_MATCH_THRESHOLD = 85


# ---------------------------------------------------------------------------
# 1. Text extraction
# ---------------------------------------------------------------------------


def extract_text_pdf(data: bytes) -> tuple[str, bool]:
    """Extract text from a PDF; OCR pages with no text layer if possible.

    Returns ``(text, used_ocr)``.
    """
    import fitz  # PyMuPDF

    used_ocr = False
    chunks: list[str] = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        for page in doc:
            text = page.get_text("text")
            if text.strip():
                chunks.append(text)
            elif _OCR_AVAILABLE:
                # Scanned/image page -> rasterize and OCR
                import pytesseract
                from PIL import Image

                pix = page.get_pixmap(dpi=200)
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                ocr_text = pytesseract.image_to_string(img)
                if ocr_text.strip():
                    chunks.append(ocr_text)
                    used_ocr = True
    return "\n".join(chunks), used_ocr


def extract_text_docx(data: bytes) -> str:
    """Extract text from a DOCX (paragraphs + tables)."""
    import docx

    document = docx.Document(io.BytesIO(data))
    lines = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            lines.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# 2. Section detection
# ---------------------------------------------------------------------------


def _match_section_header(line: str) -> str | None:
    """Return the canonical section key if the line looks like a header."""
    stripped = line.strip().strip(":").strip()
    # Headers are short standalone lines, not sentences
    if not stripped or len(stripped.split()) > 5 or len(stripped) > 45:
        return None
    lowered = stripped.lower()

    for section, aliases in _SECTION_ALIASES.items():
        for alias in aliases:
            if lowered == alias:
                return section
            if _fuzz is not None and _fuzz.ratio(lowered, alias) >= _HEADER_MATCH_THRESHOLD:
                return section
    return None


def split_sections(text: str) -> dict[str, list[str]]:
    """Split raw resume text into ``{section: [lines]}``.

    Lines before the first recognized header go to the ``_header`` block
    (name + contact info usually live there).
    """
    sections: dict[str, list[str]] = {"_header": []}
    current = "_header"

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        matched = _match_section_header(line)
        if matched:
            current = matched
            sections.setdefault(current, [])
            continue
        sections.setdefault(current, []).append(line)

    return sections


# ---------------------------------------------------------------------------
# 3. Field extraction helpers
# ---------------------------------------------------------------------------


def _normalize_date(token: str) -> str:
    """Convert a date token to ``YYYY-MM`` (month inputs) or empty string."""
    token = token.strip().lower()
    if token in ("present", "current", "now", "ongoing"):
        return ""
    year_m = re.search(r"(19|20)\d{2}", token)
    if not year_m:
        return ""
    year = year_m.group(0)
    month = "01"
    for abbr, num in _MONTH_NUM.items():
        if token.startswith(abbr):
            month = num
            break
    return f"{year}-{month}"


def _extract_contact(header_lines: list[str], full_text: str) -> dict[str, Any]:
    """Pull name/email/phone/linkedin/location from the header block."""
    header_text = "\n".join(header_lines)
    search_text = header_text or full_text

    email_m = _EMAIL_RE.search(search_text) or _EMAIL_RE.search(full_text)
    phone_m = _PHONE_RE.search(search_text) or _PHONE_RE.search(full_text)
    linkedin_m = _LINKEDIN_RE.search(search_text) or _LINKEDIN_RE.search(full_text)

    email = email_m.group(0) if email_m else ""
    phone = re.sub(r"[^\d]", "", phone_m.group(0))[-10:] if phone_m else ""
    linkedin = linkedin_m.group(0) if linkedin_m else ""
    if linkedin and not linkedin.lower().startswith("http"):
        linkedin = "https://" + linkedin

    # Name: first header line that isn't contact info / a URL, 2-5 words.
    name = ""
    for line in header_lines[:5]:
        if _EMAIL_RE.search(line) or _PHONE_RE.search(line) or _URL_RE.search(line):
            continue
        words = line.split()
        if 1 < len(words) <= 5 and not any(ch.isdigit() for ch in line):
            name = line.strip()
            break
    # spaCy PERSON entity as a cross-check / fallback
    if not name and _SPACY_NLP is not None and header_lines:
        doc = _SPACY_NLP(" ".join(header_lines[:5]))
        persons = [e.text for e in doc.ents if e.label_ == "PERSON"]
        if persons:
            name = persons[0]

    # Location: "City, ST"-style token in header (short, has a comma, no digits)
    location = ""
    for line in header_lines[:6]:
        for part in re.split(r"[|•·]+", line):
            part = part.strip()
            if (
                "," in part
                and len(part.split()) <= 4
                and not any(ch.isdigit() for ch in part)
                and not _EMAIL_RE.search(part)
                and "linkedin" not in part.lower()
            ):
                location = part
                break
        if location:
            break

    return {
        "full_name": name,
        "email": email,
        "phone": phone,
        "linkedin_url": linkedin,
        "location": location,
    }


def _spacy_orgs(text: str) -> list[str]:
    if _SPACY_NLP is None:
        return []
    doc = _SPACY_NLP(text[:500])
    return [e.text.strip() for e in doc.ents if e.label_ == "ORG"]


def _extract_education(lines: list[str]) -> list[dict[str, str]]:
    """Group education lines into entries anchored on degree keywords."""
    entries: list[dict[str, str]] = []
    current: dict[str, str] | None = None

    for line in lines:
        degree_m = _DEGREE_RE.search(line)
        date_m = _DATE_RANGE_RE.search(line)
        gpa_m = _GPA_RE.search(line)

        # A degree mention starts a new entry
        if degree_m:
            if current:
                entries.append(current)
            current = {
                "institution": "", "degree": "", "field_of_study": "",
                "start_date": "", "end_date": "", "gpa": "",
            }
            current["degree"] = degree_m.group(0).strip()
            # Text after the degree token is usually the field of study
            after = line[degree_m.end():].strip(" ,-–:in")
            if after and len(after) < 60 and not _DATE_RANGE_RE.search(after):
                current["field_of_study"] = after.strip()
            # Institution may be on the same line before the degree
            before = line[: degree_m.start()].strip(" ,-–|")
            if len(before) > 3:
                current["institution"] = before
        elif current is not None and not current["institution"]:
            # Line following a degree without institution — likely the college
            orgs = _spacy_orgs(line)
            candidate = orgs[0] if orgs else (line if len(line) < 80 else "")
            if candidate and not date_m:
                current["institution"] = candidate.strip(" ,|")

        if current is not None:
            if date_m:
                current["start_date"] = _normalize_date(date_m.group(1))
                current["end_date"] = _normalize_date(date_m.group(2))
            if gpa_m and not current["gpa"]:
                current["gpa"] = gpa_m.group(1).strip()

    if current:
        entries.append(current)
    return entries


# Bullet-line prefixes. "?" included: PDFs with unmapped glyphs often
# extract "•" as "?" — treat it as a bullet artifact, not content.
_BULLET_PREFIXES = ("•", "-", "–", "▪", "‣", "·", "*", "?")
_BULLET_STRIP = "•-–▪‣·*? "


def _looks_like_entry_heading(line: str) -> bool:
    """Heuristic: short line, not a bullet, likely 'Title @ Company' heading."""
    if line.lstrip().startswith(_BULLET_PREFIXES):
        return False
    return len(line.split()) <= 12


def _split_heading(text: str) -> tuple[str, str]:
    """Split 'Title | Company' / 'Title, Company' / 'Title - Company'."""
    parts = [p.strip() for p in re.split(r"[|,•]+|\s[-–]\s", text) if p.strip()]
    if len(parts) >= 2:
        return parts[0], parts[1]
    if parts:
        return parts[0], ""
    return "", ""


def _extract_experience(lines: list[str]) -> list[dict[str, Any]]:
    """Group experience lines into entries anchored on date ranges."""
    entries: list[dict[str, Any]] = []
    current: dict[str, Any] | None = None
    desc: list[str] = []

    def flush() -> None:
        nonlocal current, desc
        if current is not None:
            current["description"] = "\n".join(desc).strip()
            entries.append(current)
        current, desc = None, []

    for i, line in enumerate(lines):
        date_m = _DATE_RANGE_RE.search(line)
        if date_m and _looks_like_entry_heading(line):
            # Heading text: same line minus the dates, or the previous
            # heading-like line ("Title | Company" above the date line).
            rest = (line[: date_m.start()] + line[date_m.end():]).strip(" ,-–|")
            if not rest:
                if desc and _looks_like_entry_heading(desc[-1]) and not _DATE_RANGE_RE.search(desc[-1]):
                    rest = desc.pop()
                elif current is None and i > 0 and _looks_like_entry_heading(lines[i - 1]):
                    rest = lines[i - 1]
            flush()
            end_raw = date_m.group(2).lower()
            title, company = _split_heading(rest)
            current = {
                "company": company, "title": title,
                "start_date": _normalize_date(date_m.group(1)),
                "end_date": _normalize_date(date_m.group(2)),
                "is_current": end_raw in ("present", "current", "now", "ongoing"),
                "description": "",
            }
            # Company still unknown — try spaCy ORG on the previous line
            if not current["company"] and i > 0:
                prev = lines[i - 1].strip()
                if _looks_like_entry_heading(prev) and not _DATE_RANGE_RE.search(prev):
                    orgs = _spacy_orgs(prev)
                    if orgs:
                        current["company"] = orgs[0]
        elif current is not None:
            desc.append(line.lstrip(_BULLET_STRIP).strip() if line.lstrip().startswith(_BULLET_PREFIXES) else line)

    flush()

    # No date-anchored entries found: dump everything as one entry so the
    # student still gets the text to edit rather than losing it.
    if not entries and lines:
        entries.append({
            "company": "", "title": "",
            "start_date": "", "end_date": "", "is_current": False,
            "description": "\n".join(lines).strip(),
        })
    return entries


def _extract_projects(lines: list[str]) -> list[dict[str, str]]:
    """Group project lines: heading lines start entries, bullets describe them."""
    entries: list[dict[str, str]] = []
    current: dict[str, str] | None = None
    desc: list[str] = []

    def flush() -> None:
        nonlocal current, desc
        if current is not None:
            current["description"] = "\n".join(desc).strip()
            entries.append(current)
        current, desc = None, []

    for line in lines:
        is_bullet = line.lstrip().startswith(_BULLET_PREFIXES)
        link_m = _URL_RE.search(line)
        tech_m = re.match(r"(?:tech(?:nologies)?|tools|stack)\s*[:\-]\s*(.+)", line, re.I)

        if tech_m and current is not None:
            current["technologies"] = tech_m.group(1).strip()
        elif not is_bullet and len(line.split()) <= 10 and not tech_m:
            # Short non-bullet line -> new project heading
            flush()
            name = _URL_RE.sub("", line).strip(" ,-–|:")
            current = {"name": name, "description": "", "technologies": "", "link": ""}
            if link_m:
                url = link_m.group(0)
                current["link"] = url if url.lower().startswith("http") else "https://" + url
        elif current is not None:
            if link_m and not current["link"]:
                url = link_m.group(0)
                current["link"] = url if url.lower().startswith("http") else "https://" + url
            stripped = _URL_RE.sub("", line).lstrip(_BULLET_STRIP).strip() if link_m else line.lstrip(_BULLET_STRIP).strip()
            if stripped:
                desc.append(stripped)
        else:
            # Description before any heading — start an unnamed project
            current = {"name": "", "description": "", "technologies": "", "link": ""}
            desc.append(line.lstrip(_BULLET_STRIP).strip())

    flush()
    return entries


def _extract_skills(lines: list[str]) -> list[str]:
    """Split skill lines on common delimiters, drop category prefixes."""
    skills: list[str] = []
    for line in lines:
        # Drop "Languages:" style category prefixes
        line = re.sub(r"^[A-Za-z &/]+:\s*", "", line.strip())
        for token in _SKILL_SPLIT_RE.split(line):
            token = token.strip(" .•-–")
            if token and 1 < len(token) <= 40:
                skills.append(token)
    # Dedupe preserving order
    seen: set[str] = set()
    out: list[str] = []
    for s in skills:
        key = s.lower()
        if key not in seen:
            seen.add(key)
            out.append(s)
    return out


# ---------------------------------------------------------------------------
# 4. Confidence
# ---------------------------------------------------------------------------


def _confidence(data: dict[str, Any], sections: dict[str, list[str]], used_ocr: bool) -> dict[str, str]:
    """Assign "high" / "low" per top-level field for UI highlighting.

    Regex-validated fields (email/phone/linkedin) are high when present.
    List sections are high when detected via an explicit header AND every
    entry has its anchor fields; anything OCR-derived is capped at low.
    """
    conf: dict[str, str] = {}

    conf["email"] = "high" if data["email"] else "low"
    conf["phone"] = "high" if data["phone"] else "low"
    conf["linkedin_url"] = "high" if data["linkedin_url"] else "low"
    conf["full_name"] = "high" if data["full_name"] else "low"
    conf["location"] = "low"  # location heuristic is always worth a check
    conf["summary"] = "high" if "summary" in sections and data["summary"] else "low"

    edu = data["education"]
    conf["education"] = (
        "high"
        if "education" in sections and edu and all(e["institution"] and e["degree"] for e in edu)
        else "low"
    )
    exp = data["experience"]
    conf["experience"] = (
        "high"
        if "experience" in sections and exp and all(e["title"] and e["start_date"] for e in exp)
        else "low"
    )
    proj = data["projects"]
    conf["projects"] = (
        "high" if "projects" in sections and proj and all(p["name"] for p in proj) else "low"
    )
    conf["skills"] = "high" if "skills" in sections and data["skills"] else "low"
    conf["certifications"] = (
        "high" if "certifications" in sections and data["certifications"] else "low"
    )

    if used_ocr:
        conf = {k: "low" for k in conf}
    return conf


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def parse_resume(data: bytes, filename: str) -> dict[str, Any]:
    """Parse an uploaded resume file into pre-fill data + confidence map.

    Returns::

        {
            "data": {...ResumeCreate-shaped fields...},
            "confidence": {field: "high" | "low"},
            "warnings": [str, ...],
            "used_ocr": bool,
        }

    Raises ``ValueError`` for unsupported/unreadable files.
    """
    ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError("Unsupported file type. Upload a PDF or DOCX file.")
    if len(data) > MAX_FILE_SIZE:
        raise ValueError("File too large. Maximum size is 5 MB.")

    warnings: list[str] = []
    used_ocr = False

    try:
        if ext == ".pdf":
            text, used_ocr = extract_text_pdf(data)
        else:
            text = extract_text_docx(data)
    except ValueError:
        raise
    except Exception as exc:  # corrupt / password-protected file
        logger.warning("Resume parse failed for %s: %s", filename, exc)
        raise ValueError("Could not read this file. Is it a valid PDF/DOCX?") from exc

    if not text.strip():
        if ext == ".pdf" and not _OCR_AVAILABLE:
            raise ValueError(
                "No selectable text found (scanned resume?) and OCR is not "
                "available on the server."
            )
        raise ValueError("No text could be extracted from this file.")

    if used_ocr:
        warnings.append(
            "This looks like a scanned resume — text was recovered via OCR, "
            "so please verify everything carefully."
        )

    sections = split_sections(text)

    contact = _extract_contact(sections.get("_header", []), text)
    summary = " ".join(sections.get("summary", [])).strip()
    education = _extract_education(sections.get("education", []))
    experience = _extract_experience(sections.get("experience", []))
    projects = _extract_projects(sections.get("projects", []))
    skills = _extract_skills(sections.get("skills", []))
    certifications = _extract_skills(sections.get("certifications", []))

    detected = [k for k in sections if k != "_header"]
    if not detected:
        warnings.append(
            "No standard section headers were detected — the resume text was "
            "imported but may need manual reorganization."
        )

    resume_data: dict[str, Any] = {
        **contact,
        "summary": summary[:1500],
        "education": education,
        "experience": experience,
        "projects": projects,
        "skills": skills[:40],
        "certifications": certifications[:20],
    }

    missing = [
        f for f in ("full_name", "email", "phone")
        if not resume_data[f]
    ]
    if missing:
        warnings.append(
            "Could not find: " + ", ".join(m.replace("_", " ") for m in missing)
            + ". Please fill these in manually."
        )

    return {
        "data": resume_data,
        "confidence": _confidence(resume_data, sections, used_ocr),
        "warnings": warnings,
        "used_ocr": used_ocr,
        "detected_sections": detected,
    }
